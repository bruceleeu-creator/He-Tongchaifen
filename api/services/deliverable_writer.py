"""
交付成果文档生成器
基于 python-docx 生成 .docx 文件
"""
import json
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml

from config import settings


class DeliverableWriter:
    """交付成果文档写入器"""

    def __init__(self):
        self.templates_path = settings.API_DIR / "data" / "deliverable_templates.json"

    def load_templates(self) -> list[dict]:
        """加载模板库"""
        if not self.templates_path.exists():
            return []
        try:
            with open(self.templates_path, "r", encoding="utf-8") as f:
                data = json.load(f)
                return data.get("templates", [])
        except (json.JSONDecodeError, IOError):
            return []

    def find_template(self, template_key: str) -> Optional[dict]:
        """按模板键查找模板"""
        templates = self.load_templates()
        for t in templates:
            if t.get("template_key") == template_key:
                return t
        return None

    def match_template(self, service_module: str, task_type: str, task_name: str) -> Optional[dict]:
        """根据任务信息匹配最佳模板"""
        templates = self.load_templates()
        # 1. 完全匹配 service_module + task_type
        for t in templates:
            if t.get("applicable_service_module") == service_module and t.get("applicable_task_type") == task_type:
                return t
        # 2. 匹配 task_type
        for t in templates:
            if t.get("applicable_task_type") == task_type:
                return t
        # 3. 匹配 service_module
        for t in templates:
            if t.get("applicable_service_module") == service_module:
                return t
        # 4. 关键词匹配
        name_lower = task_name.lower()
        keyword_map = {
            "启动会": "startup_meeting_minutes",
            "资料": "client_document_list",
            "清单": "client_document_list",
            "确认": "milestone_confirmation",
            "验收": "milestone_confirmation",
            "月报": "monthly_report",
            "周报": "monthly_report",
            "报告": "monthly_report",
            "延期": "delay_notice",
            "顺延": "delay_notice",
        }
        for keyword, key in keyword_map.items():
            if keyword in name_lower:
                return self.find_template(key)
        return None

    def _set_cell_border(self, cell, **kwargs):
        """设置表格单元格边框"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            tcPr.append(tcBorders)
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = parse_xml(r'<w:{} xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'.format(edge))
                    tcBorders.append(element)
                for key in ["sz", "val", "color", "space"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

    def _fill_variables(self, text: str, variables: dict) -> str:
        """替换文本中的变量标记"""
        result = text
        for key, value in variables.items():
            placeholder = "{" + key + "}"
            result = result.replace(placeholder, str(value) if value else "")
        return result

    def generate_docx(self, artifact: dict, output_path: Path) -> Path:
        """生成 docx 文件

        Args:
            artifact: 交付成果数据
            output_path: 输出文件路径

        Returns:
            生成的文件路径
        """
        doc = Document()

        # 设置默认中文字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(10.5)

        # 获取任务变量
        variables = artifact.get("variables", {})
        template_key = artifact.get("template_key", "")
        template = self.find_template(template_key) if template_key else None

        # 标题
        deliverable_name = artifact.get("deliverable_name", "交付成果")
        title_text = self._fill_variables(deliverable_name, variables)
        title = doc.add_heading(title_text, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1B, 0x23, 0x32)

        # 副标题：客户和项目信息
        customer = variables.get("customer_name", "")
        project = variables.get("project_name", "")
        if customer or project:
            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle.add_run(f"客户：{customer}    项目：{project}")
            subtitle_run.font.size = Pt(10)
            subtitle_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # 分隔线
        doc.add_paragraph()

        # 优先使用 content_sections（内容丰满的章节，来自 deliverable_design 规则模式）
        content_sections = artifact.get("content_sections", [])
        if content_sections:
            # AI 设计说明
            ai_reason = artifact.get("ai_design_reason", "")
            if ai_reason:
                h = doc.add_heading("成果设计说明", level=2)
                for run in h.runs:
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0x1B, 0x23, 0x32)
                p = doc.add_paragraph(ai_reason)
                for run in p.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                doc.add_paragraph()

            # 逐章节写入标题 + 正文要点（来自任务实际字段，无空段）
            for section in content_sections:
                section_title = section.get("title", "")
                bullets = section.get("bullets", [])
                heading = doc.add_heading(section_title, level=2)
                for run in heading.runs:
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0x1B, 0x23, 0x32)
                for bullet in bullets:
                    p = doc.add_paragraph(style='List Bullet')
                    run = p.add_run(str(bullet))
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10.5)
                doc.add_paragraph()

            # 验收标准
            acceptance_criteria = artifact.get("acceptance_criteria", [])
            if acceptance_criteria:
                h = doc.add_heading("验收标准", level=2)
                for run in h.runs:
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0x1B, 0x23, 0x32)
                for item in acceptance_criteria:
                    p = doc.add_paragraph(style='List Bullet')
                    run = p.add_run(str(item))
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10.5)
                doc.add_paragraph()

            # 客户需提供的资料
            client_inputs = artifact.get("client_inputs", [])
            if client_inputs:
                h = doc.add_heading("客户需提供的资料", level=2)
                for run in h.runs:
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0x1B, 0x23, 0x32)
                for item in client_inputs:
                    p = doc.add_paragraph(style='List Bullet')
                    run = p.add_run(str(item))
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10.5)
                doc.add_paragraph()

            # 风险提示
            risk_notes = artifact.get("risk_notes", [])
            if risk_notes:
                h = doc.add_heading("风险提示", level=2)
                for run in h.runs:
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0x1B, 0x23, 0x32)
                for item in risk_notes:
                    p = doc.add_paragraph(style='List Bullet')
                    run = p.add_run(str(item))
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10.5)
                doc.add_paragraph()

            # 下一步动作
            next_actions = artifact.get("next_actions", [])
            if next_actions:
                h = doc.add_heading("下一步动作", level=2)
                for run in h.runs:
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0x1B, 0x23, 0x32)
                for item in next_actions:
                    p = doc.add_paragraph(style='List Number')
                    run = p.add_run(str(item))
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10.5)
                doc.add_paragraph()

            # 关联任务信息表格
            h = doc.add_heading("关联任务信息", level=2)
            for run in h.runs:
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x1B, 0x23, 0x32)

            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '字段'
            hdr_cells[1].text = '内容'
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = '黑体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

            task_fields = [
                ("任务名称", variables.get("task_name", "")),
                ("服务模块", variables.get("service_module", "")),
                ("任务类型", variables.get("task_type", "")),
                ("里程碑目标", variables.get("milestone_goal", "")),
                ("交付成果", variables.get("deliverables", "")),
                ("计划完成时间", variables.get("plan_end_date", "")),
                ("我方负责人", variables.get("our_owner", "")),
                ("客户责任人", variables.get("client_contact", "")),
            ]
            for field_name, field_value in task_fields:
                row_cells = table.add_row().cells
                row_cells[0].text = field_name
                row_cells[1].text = str(field_value) if field_value else ""
        elif template:
            # 兼容：有模板但无 content_sections 时按模板生成（输出具体描述）
            content_schema = template.get("content_schema", {})
            sections = content_schema.get("sections", [])
            for section in sections:
                section_text = self._fill_variables(section, variables)
                heading = doc.add_heading(section_text, level=2)
                for run in heading.runs:
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0x1B, 0x23, 0x32)
                p = doc.add_paragraph()
                run = p.add_run(f"本章节围绕{variables.get('task_name', '本任务')}展开，涉及{variables.get('service_module', '相关服务')}模块，依据合同约定与项目计划执行。")
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                doc.add_paragraph()
        else:
            # 无模板时，按通用结构生成
            # AI 设计说明
            ai_reason = artifact.get("ai_design_reason", "")
            if ai_reason:
                h = doc.add_heading("成果设计说明", level=2)
                for run in h.runs:
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                p = doc.add_paragraph(ai_reason)
                for run in p.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                doc.add_paragraph()

            # 内容大纲
            outline = artifact.get("content_outline", [])
            if outline:
                h = doc.add_heading("内容大纲", level=2)
                for run in h.runs:
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                for item in outline:
                    p = doc.add_paragraph(style='List Number')
                    p.add_run(str(item))
                doc.add_paragraph()

            # 任务信息表格
            h = doc.add_heading("关联任务信息", level=2)
            for run in h.runs:
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '字段'
            hdr_cells[1].text = '内容'
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = '黑体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

            task_fields = [
                ("任务名称", variables.get("task_name", "")),
                ("服务模块", variables.get("service_module", "")),
                ("任务类型", variables.get("task_type", "")),
                ("计划完成时间", variables.get("plan_end_date", "")),
                ("我方负责人", variables.get("our_owner", "")),
                ("客户责任人", variables.get("client_contact", "")),
            ]
            for field_name, field_value in task_fields:
                row_cells = table.add_row().cells
                row_cells[0].text = field_name
                row_cells[1].text = str(field_value) if field_value else ""

        # 页脚：生成信息
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}    编号：{artifact.get('artifact_id', '')}")
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # 保存
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return output_path


# 全局实例
deliverable_writer = DeliverableWriter()
