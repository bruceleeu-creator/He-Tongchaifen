"""
Word 文档解析服务
使用 python-docx 解析上传的 .docx 文件
"""
from pathlib import Path
from typing import Optional
import re
import uuid
import shutil

from config import settings


class DocxParser:
    """Word 文档解析器"""

    def __init__(self):
        self.upload_dir = settings.UPLOAD_DIR
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    def save_upload(self, file_content: bytes, filename: str, run_id: str) -> str:
        """保存上传文件到 run 目录下的 uploads 子目录

        Returns:
            保存后的文件路径
        """
        upload_dir = settings.get_run_dir(run_id) / "uploads"
        upload_dir.mkdir(parents=True, exist_ok=True)

        # 生成唯一文件名（扩展名白名单校验，防路径穿越）
        ext = Path(filename).suffix
        if not re.match(r"^\.[A-Za-z0-9]{1,10}$", ext or ""):
            ext = ".docx"
        saved_filename = f"{uuid.uuid4().hex[:12]}{ext}"
        filepath = upload_dir / saved_filename
        filepath.write_bytes(file_content)

        return str(filepath)

    def parse(self, filepath: str) -> dict:
        """解析 Word 文档，提取文本内容

        Returns:
            {
                "filename": str,
                "paragraphs": list[str],
                "tables": list[list[list[str]]],
                "full_text": str,
            }
        """
        filepath = Path(filepath)
        if not filepath.exists():
            raise FileNotFoundError(f"文件不存在: {filepath}")

        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx 未安装，请运行: pip install python-docx")

        doc = Document(str(filepath))

        # 提取段落
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 提取表格
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        full_text = "\n".join(paragraphs)

        # 如果有表格，也把表格内容加入 full_text
        if tables:
            table_texts = []
            for i, table in enumerate(tables):
                table_texts.append(f"\n--- 表格 {i + 1} ---")
                for row in table:
                    table_texts.append(" | ".join(row))
            full_text += "\n" + "\n".join(table_texts)

        return {
            "filename": filepath.name,
            "paragraphs": paragraphs,
            "tables": tables,
            "full_text": full_text,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "char_count": len(full_text),
        }

    def parse_to_text(self, filepath: str) -> str:
        """解析 Word 文档为纯文本"""
        result = self.parse(filepath)
        return result["full_text"]


# 全局实例
docx_parser = DocxParser()
